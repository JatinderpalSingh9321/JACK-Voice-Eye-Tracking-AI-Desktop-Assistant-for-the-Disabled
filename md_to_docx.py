import os
from docx import Document
from docx.shared import Pt
import re

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            
        # Horizontal Rule
        elif line == '---':
            doc.add_page_break() # Or just skip
            
        # Bullet points
        elif line.startswith('* ') or line.startswith('- '):
            text = line[2:]
            # Basic bold formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
            
        # Numbered lists
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Number')
            
        # Bold/Text handling
        else:
            # Simple bold removal/handling (python-docx handles inline better with runs)
            p = doc.add_paragraph()
            
            # Split by bold markers
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == "__main__":
    md_file = "NAVTOOLS_EOG_PRESENTATION.md"
    docx_file = "NAVTOOLS_EOG_PRESENTATION.docx"
    convert_md_to_docx(md_file, docx_file)
